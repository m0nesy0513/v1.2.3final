from docx import Document


def parse_docx(file_or_path) -> str:
    doc = Document(file_or_path)
    texts = []

    for p in doc.paragraphs:
        value = p.text.strip()
        if value:
            texts.append(value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    texts.append(value)

    return "\n".join(texts)
